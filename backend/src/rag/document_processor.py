"""Document processing utilities for the AI-Enhanced Interactive Book Agent.

This module provides utilities for processing various document formats
(PDF, DOCX, EPUB, TXT) and extracting text content for the RAG system.
"""
import asyncio
import pdfplumber
import docx
from ebooklib import ebooklib
from ebooklib.epub import EpubReader
import html2text
from typing import List, Tuple, Optional
from pathlib import Path


class DocumentProcessor:
    """Utility class for processing different document formats."""
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """Extract text content from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            text_content = ""
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            return text_content
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    @staticmethod
    async def extract_text_from_epub(file_path: str) -> str:
        """Extract text content from an EPUB file.
        
        Args:
            file_path: Path to the EPUB file
            
        Returns:
            Extracted text content
        """
        try:
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.ignore_images = True
            
            reader = EpubReader(file_path)
            book = reader.load_book()
            
            text_content = []
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    # Convert HTML content to plain text
                    content = h.handle(item.get_content().decode('utf-8'))
                    text_content.append(content)
            
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Error extracting text from EPUB: {str(e)}")
    
    @staticmethod
    async def extract_text_from_txt(file_path: str) -> str:
        """Extract text content from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error extracting text from TXT: {str(e)}")
    
    @staticmethod
    async def get_document_metadata(file_path: str) -> dict:
        """Extract metadata from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata
        """
        path = Path(file_path)
        metadata = {
            'file_name': path.name,
            'file_size': path.stat().st_size,
            'file_type': path.suffix.lower(),
            'last_modified': path.stat().st_mtime
        }
        
        # Add format-specific metadata
        if path.suffix.lower() == '.pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata['page_count'] = len(pdf.pages)
                    # Extract PDF-specific metadata if available
                    if pdf.metadata:
                        metadata.update({
                            'title': pdf.metadata.get('Title', ''),
                            'author': pdf.metadata.get('Author', ''),
                            'subject': pdf.metadata.get('Subject', ''),
                            'creator': pdf.metadata.get('Creator', ''),
                            'producer': pdf.metadata.get('Producer', ''),
                            'creation_date': pdf.metadata.get('CreationDate', ''),
                            'modification_date': pdf.metadata.get('ModDate', ''),
                        })
            except:
                pass  # If we can't get PDF metadata, continue with basic metadata
        
        elif path.suffix.lower() == '.docx':
            try:
                doc = docx.Document(file_path)
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'creator': core_props.creator or '',
                    'description': core_props.description or '',
                    'keywords': core_props.keywords or '',
                    'last_modified_by': core_props.last_modified_by or '',
                    'revision': core_props.revision,
                    'created': core_props.created,
                    'modified': core_props.modified,
                })
            except:
                pass  # If we can't get DOCX metadata, continue with basic metadata
        
        return metadata
    
    @staticmethod
    async def process_document(file_path: str, chunk_size: int = 1000, overlap: int = 100) -> List[Tuple[str, dict]]:
        """Process a document and split it into chunks for the RAG system.
        
        Args:
            file_path: Path to the document file
            chunk_size: Size of each text chunk
            overlap: Overlap between chunks to maintain context
            
        Returns:
            List of tuples containing (chunk_text, chunk_metadata)
        """
        # Determine file type and extract text
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            text_content = await DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            text_content = await DocumentProcessor.extract_text_from_docx(file_path)
        elif file_ext == '.epub':
            text_content = await DocumentProcessor.extract_text_from_epub(file_path)
        elif file_ext == '.txt':
            text_content = await DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Split text into chunks
        chunks = DocumentProcessor._split_text(text_content, chunk_size, overlap)
        
        # Get document metadata
        doc_metadata = await DocumentProcessor.get_document_metadata(file_path)
        
        # Create chunks with metadata
        chunked_data = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = doc_metadata.copy()
            chunk_metadata.update({
                'chunk_id': f"{doc_metadata['file_name']}_chunk_{i}",
                'chunk_index': i,
                'total_chunks': len(chunks),
                'start_pos': i * (chunk_size - overlap),
                'end_pos': min((i + 1) * (chunk_size - overlap) + chunk_size, len(text_content))
            })
            chunked_data.append((chunk, chunk_metadata))
        
        return chunked_data
    
    @staticmethod
    def _split_text(text: str, chunk_size: int, overlap: int) -> List[str]:
        """Split text into chunks of specified size with overlap.
        
        Args:
            text: Text to be split
            chunk_size: Size of each text chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # If this is not the last chunk, try to break at a sentence boundary
            if end < len(text):
                # Look for a sentence boundary near the end
                search_start = end - overlap
                sentence_end = -1
                
                for sep in ['.\n', '. ', '! ', '? ', '\n\n', '\n']:
                    last_sep = text.rfind(sep, search_start, end)
                    if last_sep != -1:
                        sentence_end = last_sep + len(sep)
                        break
                
                if sentence_end != -1 and sentence_end > search_start:
                    end = sentence_end
                else:
                    # If no good break point found, just take the chunk
                    end = min(end, len(text))
            
            chunk = text[start:end].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            start = end  # Move start to the end of current chunk
            
            # If overlap is required, adjust start position
            if overlap > 0 and start > overlap:
                start = start - overlap
        
        return chunks


# Global instance of the document processor
document_processor = DocumentProcessor()