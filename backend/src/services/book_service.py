"""Service layer for book operations in AI-Enhanced Interactive Book Agent."""
from typing import List, Optional
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy import update, delete
import uuid
import os
from pathlib import Path
from fastapi import UploadFile

from backend.src.models.sqlalchemy_models import Book as SQLAlchemyBook, User
from backend.src.models.book import Book as PydanticBook
from backend.src.config import settings


class BookService:
    """Service class for handling book operations."""
    
    def __init__(self, db: AsyncSession):
        """Initialize the service with a database session."""
        self.db = db

    async def create_book_record(self, user_id: str, file: UploadFile, file_extension: str) -> PydanticBook:
        """Create a record for an uploaded book."""
        # Verify the user exists
        user_exists = await self.db.execute(
            select(User).filter(User.id == user_id)
        )
        user = user_exists.scalars().first()
        if not user:
            raise ValueError(f"User with id {user_id} does not exist")
        
        # Sanitize filename
        filename = f"{uuid.uuid4()}_{file.filename}"
        
        # Define upload path
        upload_dir = Path("uploads") / str(user_id)
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / filename
        
        # Save the file
        with open(file_path, "wb") as buffer:
            buffer.write(file.file.read())
        
        # Create the book record
        db_book = SQLAlchemyBook(
            user_id=user_id,
            title=file.filename.replace(f".{file_extension}", ""),
            author="Unknown",  # Will be extracted during processing
            file_path=str(file_path),
            file_format=file_extension,
            file_size=file.file.tell() if hasattr(file.file, 'tell') else 0,
            total_pages=0  # Will be determined during processing
        )
        
        self.db.add(db_book)
        await self.db.commit()
        await self.db.refresh(db_book)
        
        # Convert to Pydantic model and return
        return self._convert_sqlalchemy_to_pydantic(db_book)

    async def get_user_books(self, user_id: str) -> List[PydanticBook]:
        """Get all books for a specific user."""
        result = await self.db.execute(
            select(SQLAlchemyBook).filter(SQLAlchemyBook.user_id == user_id)
        )
        books = result.scalars().all()
        
        return [self._convert_sqlalchemy_to_pydantic(book) for book in books]

    async def get_book_by_id(self, book_id: str) -> Optional[PydanticBook]:
        """Get a specific book by its ID."""
        try:
            # Validate UUID format
            uuid.UUID(book_id)
        except ValueError:
            return None
            
        result = await self.db.execute(
            select(SQLAlchemyBook).filter(SQLAlchemyBook.id == book_id)
        )
        book = result.scalars().first()
        
        if book:
            return self._convert_sqlalchemy_to_pydantic(book)
        return None

    async def delete_book(self, book_id: str) -> bool:
        """Delete a book record and its associated file."""
        try:
            # Validate UUID format
            uuid.UUID(book_id)
        except ValueError:
            raise ValueError(f"Invalid book ID: {book_id}")
        
        # Get the book to retrieve its file path
        result = await self.db.execute(
            select(SQLAlchemyBook).filter(SQLAlchemyBook.id == book_id)
        )
        book = result.scalars().first()
        
        if not book:
            return False
        
        # Delete the file from the filesystem
        try:
            if os.path.exists(book.file_path):
                os.remove(book.file_path)
        except OSError:
            # If file deletion fails, log the error but continue with DB deletion
            print(f"Warning: Could not delete file {book.file_path}")
        
        # Delete the book record from the database
        stmt = delete(SQLAlchemyBook).where(SQLAlchemyBook.id == book_id)
        await self.db.execute(stmt)
        await self.db.commit()
        
        return True

    async def process_book_content(self, book_id: str) -> bool:
        """Start processing a book to extract content and metadata."""
        from backend.src.rag.rag_engine import RAGEngine
        from backend.src.models.sqlalchemy_models import BookContent
        import pdfplumber
        from docx import Document
        import ebooklib
        from ebooklib import epub

        # This method would typically start a background task
        # For now, we'll implement the actual processing
        try:
            # Validate UUID format
            uuid.UUID(book_id)
        except ValueError:
            raise ValueError(f"Invalid book ID: {book_id}")

        # Get the book record
        result = await self.db.execute(
            select(SQLAlchemyBook).filter(SQLAlchemyBook.id == book_id)
        )
        book = result.scalars().first()

        if not book:
            return False

        # Update the book's processing status
        stmt = (
            update(SQLAlchemyBook)
            .where(SQLAlchemyBook.id == book_id)
            .values(is_processed=False, processing_error=None)  # Mark as processing
        )

        result = await self.db.execute(stmt)
        await self.db.commit()

        if result.rowcount == 0:
            return False

        try:
            # Extract text and metadata based on file type
            content = ""
            total_pages = 0

            if book.file_format == 'pdf':
                with pdfplumber.open(book.file_path) as pdf:
                    content = '\n'.join([page.extract_text() for page in pdf.pages if page.extract_text()])
                    total_pages = len(pdf.pages)
            elif book.file_format == 'docx':
                doc = Document(book.file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                total_pages = len(doc.paragraphs) // 10  # Rough approximation
            elif book.file_format == 'epub':
                book_obj = epub.read_epub(book.file_path)
                content = ''
                for item in book_obj.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        content += item.get_content().decode('utf-8')
            elif book.file_format == 'txt':
                with open(book.file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            # Update the book with extracted metadata
            stmt = (
                update(SQLAlchemyBook)
                .where(SQLAlchemyBook.id == book_id)
                .values(
                    total_pages=total_pages,
                    is_processed=True
                )
            )
            await self.db.execute(stmt)
            await self.db.commit()

            # Split content into chunks
            chunk_size = 1000  # Characters per chunk
            content_chunks = []
            for i in range(0, len(content), chunk_size):
                chunk = content[i:i + chunk_size]
                content_chunks.append({
                    'content': chunk,
                    'chunk_id': f"{book_id}_chunk_{i // chunk_size}",
                    'page_number': (i // chunk_size) + 1,
                    'section_title': f'Page {(i // chunk_size) + 1}',
                    'book_id': book_id
                })

            # Store content chunks in the database
            for chunk_data in content_chunks:
                content_chunk = BookContent(
                    book_id=chunk_data['book_id'],
                    chunk_id=chunk_data['chunk_id'],
                    content=chunk_data['content'],
                    page_number=chunk_data['page_number'],
                    section_title=chunk_data['section_title'],
                    chapter=f"Chapter {(chunk_data['page_number'] - 1) // 10 + 1}"  # Group every 10 pages into a chapter
                )
                self.db.add(content_chunk)

            await self.db.commit()

            # Index the content in the RAG system
            rag_engine = RAGEngine()
            await rag_engine.add_book_content_to_index(self.db, book_id)

            return True

        except Exception as e:
            # If processing fails, log the error and update the book record
            error_msg = str(e)
            stmt = (
                update(SQLAlchemyBook)
                .where(SQLAlchemyBook.id == book_id)
                .values(
                    is_processed=False,
                    processing_error=error_msg
                )
            )
            await self.db.execute(stmt)
            await self.db.commit()
            print(f"Error processing book {book_id}: {error_msg}")
            return False

    def _convert_sqlalchemy_to_pydantic(self, db_book: SQLAlchemyBook) -> PydanticBook:
        """
        Convert SQLAlchemy Book model to Pydantic Book model.
        """
        return PydanticBook(
            id=str(db_book.id),
            user_id=str(db_book.user_id),
            title=db_book.title,
            author=db_book.author,
            file_path=db_book.file_path,
            file_format=db_book.file_format,
            file_size=db_book.file_size,
            total_pages=db_book.total_pages,
            is_processed=db_book.is_processed,
            processing_error=db_book.processing_error,
            created_at=db_book.created_at,
            updated_at=db_book.updated_at
        )